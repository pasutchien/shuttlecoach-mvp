# Generates the backend-integration handoff Word document for Chien.
# Run:  python scripts/make_handoff_docx.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x0A, 0x16, 0x28)
BLUE = RGBColor(0x25, 0x63, 0xEB)
ORANGE = RGBColor(0xE8, 0x4A, 0x30)
SLATE = RGBColor(0x47, 0x55, 0x69)

doc = Document()

# ----- base styles -----------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def code(text: str):
    """Add a monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x3A)
    return p


def body(text: str, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(text: str, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h


def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = BLUE
    return h


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = hd
        for para in cell.paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ============================ TITLE =========================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Shuttle Coach — MVP")
tr.bold = True
tr.font.size = Pt(26)
tr.font.color.rgb = BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Front-End → Backend Integration Handoff")
sr.font.size = Pt(15)
sr.font.color.rgb = NAVY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run(
    "Prepared for: Chien (Backend / AI analysis)\n"
    "AI-powered badminton form analysis  ·  Presented by Banthongyord\n"
    "Document date: 22 May 2026"
)
mr.font.size = Pt(10)
mr.font.color.rgb = SLATE

doc.add_paragraph()
intro = doc.add_paragraph()
ir = intro.add_run(
    "Hi Chien — the front-end MVP is complete. All 14 screens are built and "
    "fully clickable today on a mock data layer. This document is everything you "
    "need to plug in the real CV/AI backend. The golden rule: the entire app "
    "talks to the backend through ONE TypeScript interface. You implement that "
    "interface; nothing in the UI changes."
)
ir.italic = True

# ============================ 1. OVERVIEW ===================================
h1("1. What you are receiving")
body(
    "A React Native (Expo SDK 54) app written in TypeScript (strict). It covers "
    "the full product spec (see SPEC.md in the repo) — onboarding, the "
    "upload/trim/calibration flow, AI processing, the side-by-side analysis "
    "screen, wallet/credits, history and profile."
)
bullet(" Expo SDK 54, React Native 0.81, React 19, expo-router (file-based nav).",
       "Stack:")
bullet(" NativeWind v4 (Tailwind), Zustand state, i18n-js (Thai default / English).",
       "UI:")
bullet(" runs in the browser (react-native-web) and on iOS/Android via Expo Go.",
       "Targets:")
bullet(" today the app runs with NO backend — a mock service layer with "
       "in-memory fixtures, realistic delays and a simulated 30–90s analysis job.",
       "Data:")
body(
    "Your job is to replace the mock with the real backend. Because every backend "
    "call is funneled through one interface, this is a contained, well-defined task.",
    italic=True,
)

# ============================ 2. RUN IT =====================================
h1("2. Run the app locally (5 minutes)")
body("Prerequisite: Node.js 20+.")
code("npm install\n"
     "npm run web        # opens http://localhost:8081 (view at phone width)\n"
     "npm start          # QR code for the Expo Go app on a real device\n\n"
     "npm run typecheck  # tsc --noEmit (strict)\n"
     "npm run build:web  # static web export to dist/")
body(
    "Tip: the seeded build starts as a returning user with history. Use "
    "Profile → Sign Out to drop to a brand-new user and walk the full "
    "Splash → Onboarding → first-analysis path.",
    italic=True,
)

# ============================ 3. THE SEAM ===================================
h1("3. The integration model — one interface")
body(
    "The UI NEVER calls fetch directly. It imports `api` from "
    "`src/services/index.ts`, and `api` is an object that satisfies the "
    "`ShuttleCoachApi` interface defined in `src/services/api.ts`."
)
code("src/services/\n"
     "  api.ts            <- THE CONTRACT: the ShuttleCoachApi interface\n"
     "  index.ts          <- picks the mock or the real impl from an env flag\n"
     "  mock/             <- the current mock implementation (fixtures, job sim)\n"
     "  real/realApi.ts   <- STUB. THIS IS YOUR FILE TO IMPLEMENT.")
body("The switch is purely an environment variable:")
code('EXPO_PUBLIC_USE_MOCK_API=true    # default -> mock implementation\n'
     'EXPO_PUBLIC_USE_MOCK_API=false   # -> real implementation (your code)\n'
     'EXPO_PUBLIC_API_BASE_URL=https://api.shuttlecoach...   # your backend URL')

h2("Integration checklist")
for i, step in enumerate([
    "Open `src/services/real/realApi.ts`. Every method currently throws "
    "`NOT_IMPLEMENTED` — replace each body with a real HTTP call.",
    "`fetch` is allowed HERE and only here. Keep all network code inside this file.",
    "Match the request/response shapes exactly (Section 5). They are the same "
    "TypeScript types the UI consumes — import them from `src/types`.",
    "Map transport/HTTP failures onto the `ApiError` class with the documented "
    "error codes (Section 6).",
    "Set `EXPO_PUBLIC_USE_MOCK_API=false` and `EXPO_PUBLIC_API_BASE_URL`.",
    "Run `npm run typecheck` and click through the app. No UI or store changes "
    "should be required.",
]):
    p = doc.add_paragraph(style="List Number")
    p.add_run(step)

# ============================ 4. ANALYSIS LIFECYCLE =========================
h1("4. The analysis job lifecycle (your core area)")
body(
    "This is the heart of your integration — the CV/AI pipeline. It is "
    "asynchronous: submit → poll → fetch result."
)
body("Step 1 — submitAnalysis(input)", bold=True)
body(
    "The user has trimmed a clip, picked a stroke, and dragged 4 court-corner "
    "pins. The app calls submitAnalysis with everything the pipeline needs. "
    "On submit you deduct 100 credits and return a jobId."
)
code("submitAnalysis(input: AnalysisRequest): Promise<{ jobId: string }>\n\n"
     "AnalysisRequest {\n"
     "  clipRef: string;            // upload reference / URI of the trimmed clip\n"
     "  trimStartSec: number;\n"
     "  trimEndSec: number;         // trimEnd - trimStart is 1..15 seconds\n"
     "  strokeType: 'Smash'|'Drop_Shot'|'Clear'|'Drive'|'Net_Kill';\n"
     "  courtCorners: {             // 4 corners, each {x,y} normalised 0..1\n"
     "    tl: {x,y}; tr: {x,y}; bl: {x,y}; br: {x,y};\n"
     "  };\n"
     "  userProfileSnapshot: { heightCm: number; favoriteProId?: string };\n"
     "}")
body("Step 2 — getAnalysisStatus(jobId), polled ~every 1s by the S8 screen",
     bold=True)
code("getAnalysisStatus(jobId): Promise<AnalysisStatusResult>\n\n"
     "AnalysisStatusResult {\n"
     "  status: 'queued' | 'processing' | 'done' | 'failed';\n"
     "  progress: number;      // 0..1  (drives the progress ring)\n"
     "  etaSeconds: number;    // estimated seconds remaining\n"
     "  error?: { code, messageKey, refundable };  // only when status='failed'\n"
     "}")
body("Step 3 — getAnalysisResult(jobId), called once status is 'done'",
     bold=True)
code("getAnalysisResult(jobId): Promise<Analysis>\n\n"
     "Analysis {\n"
     "  id, jobId, createdAt,\n"
     "  strokeType,\n"
     "  overallScore: number,                  // 0..100, weighted average\n"
     "  checkpoints: { key, score }[],         // the 7 sub-scores below\n"
     "  proPlayerId: string,                   // matched pro for comparison\n"
     "  userVideoUrl: string,                  // S9 left panel\n"
     "  proVideoUrl: string,                   // S9 right panel\n"
     "  mistakes: MistakeCard[],               // <=5, sorted Critical->Minor\n"
     "  durationSec: number,\n"
     "}")
body(
    "The 7 scoring checkpoints (SPEC §8.1). Each gets a 0–100 sub-score; "
    "the overall score is their weighted average (weights vary by stroke — "
    "see `src/constants/strokes.ts`):"
)
table(
    ["key", "Measures"],
    [
        ["grip", "Finger placement / wrist angle"],
        ["stance", "Hip & shoulder rotation vs. net"],
        ["backswing", "Depth the racket arm is pulled back"],
        ["elbow", "Elbow position relative to shoulder"],
        ["contact", "Racket–shuttle contact height"],
        ["weight", "Weight transfer back foot → front"],
        ["follow_through", "Whether the swing is completed"],
    ],
    widths=[1.4, 4.6],
)
body("Each MistakeCard (max 5, plain-language coaching feedback):", bold=True)
code("MistakeCard {\n"
     "  id, title,            // short plain-language title\n"
     "  description,          // 1-2 sentences: the mistake + its impact\n"
     "  severity: 'Critical' | 'Major' | 'Minor',\n"
     "  timestampSec: number, // worst-frame time -> the app jumps the video here\n"
     "  drillId: string,      // references a drill for the 'How to Fix' sheet\n"
     "}")

# ============================ 5. FULL METHOD LIST ===========================
h1("5. Full method list (ShuttleCoachApi)")
body("All methods are async. The complete spec with suggested HTTP routes is in "
     "BACKEND.md in the repo.")
table(
    ["Method", "Purpose"],
    [
        ["getProPlayers()", "Pro player roster + reference-clip coverage"],
        ["getUserProfile()", "Current profile, or null if not onboarded"],
        ["updateUserProfile(patch)",
         "Create or patch profile. FIRST call also grants 100 free credits"],
        ["getCreditBalance()", "Current credit balance"],
        ["getTransactions()", "Credit ledger, newest first"],
        ["getCreditPackages()", "The 3 purchasable packages"],
        ["purchaseCredits(packageId)", "Process payment, add credits"],
        ["refundCredits(jobId, reason)",
         "Refund 100 credits for a failed/cancelled job"],
        ["submitAnalysis(input)", "Start a CV job; deduct 100 credits"],
        ["getAnalysisStatus(jobId)", "Poll job status / progress / ETA"],
        ["getAnalysisResult(jobId)", "Fetch the finished Analysis; save to history"],
        ["listAnalyses()", "All saved analyses, newest first"],
        ["getAnalysis(id)", "One analysis by id"],
        ["deleteAnalysis(id)", "Delete an analysis"],
        ["getDrill(drillId)", "Drill detail for the 'How to Fix' sheet"],
    ],
    widths=[2.4, 3.6],
)

# ============================ 6. ERROR CODES ================================
h1("6. Error handling")
body(
    "Throw the `ApiError` class (from `src/services/api.ts`) with a stable "
    "`code`. The UI branches on `code`, never on the message string."
)
body("Thrown ApiError codes:", bold=True)
table(
    ["code", "When"],
    [
        ["INSUFFICIENT_CREDITS", "submitAnalysis with balance < 100"],
        ["NOT_FOUND", "Unknown id (job / analysis / drill / package)"],
        ["JOB_NOT_DONE", "getAnalysisResult called before the job finished"],
        ["NETWORK_ERROR", "Transport failure"],
        ["NOT_IMPLEMENTED", "Real method not yet wired (dev only)"],
        ["UNKNOWN", "Anything else"],
    ],
    widths=[2.2, 3.8],
)
body("Analysis-job failure codes (returned INSIDE the status object, "
     "not thrown):", bold=True)
table(
    ["error.code", "Meaning"],
    [
        ["VIDEO_TOO_DARK", "AI cannot detect joints"],
        ["PLAYER_NOT_DETECTED", "No person/racket in frame"],
        ["VIDEO_TOO_BLURRY", "Frame quality below threshold"],
        ["PROCESSING_TIMEOUT", "Processing took too long"],
        ["NETWORK_ERROR / UNKNOWN", "Other failures"],
    ],
    widths=[2.2, 3.8],
)
body(
    "IMPORTANT RULE (SPEC §6.3): every job failure AFTER the credit "
    "deduction must be refundable. Set error.refundable = true; the app then "
    "calls refundCredits and shows the user a '+100 Credits refunded' toast "
    "automatically. You do not trigger the refund — the app does — but "
    "your status response must mark it refundable.",
    bold=True,
)

# ============================ 7. GOTCHAS ====================================
h1("7. Things to know before you start")
bullet(" app/login.tsx is a MOCK login screen — email/password fields plus "
       "'Continue with Google' and 'Continue with LINE'. Any input or button "
       "just proceeds; there is NO real authentication, password storage or "
       "OAuth yet. Real auth is yours to build: add the endpoints, issue a "
       "session token, and wire the screen's proceed() to a real sign-in call. "
       "The 'Sign up' link is also a placeholder. The app flow is "
       "Splash -> Login -> Onboarding/Home.",
       "Login / auth (NOT built):")
bullet(" the mock ships generic placeholder clips in assets/videos/. The real "
       "backend should return real per-analysis userVideoUrl / proVideoUrl. The "
       "app resolves them via resolveVideoSource() in src/constants/media.ts — "
       "plain URLs fall straight through.", "Sample videos:")
bullet(" MistakeCard titles/descriptions and Drill text are analysis CONTENT, "
       "not UI labels — the app renders them as-is. For a Thai-first product, "
       "return them in the user's language (send an Accept-Language header or "
       "read a profile locale field).", "Localisation:")
bullet(" updateUserProfile, on its FIRST call (no profile yet), must create the "
       "profile AND grant 100 free credits with a 'Welcome Gift' transaction "
       "(SPEC §1.5).", "Onboarding gift:")
bullet(" the mock's reset (Sign Out / Delete Account) currently goes through a "
       "mock-only `resetDb` helper. For the real backend, treat sign-out and "
       "account-deletion as real endpoints — this is flagged in "
       "IMPROVEMENTS.md under 'Needs your approval' as a small interface tweak "
       "(adding signOut()/deleteAccount() to ShuttleCoachApi). Worth doing "
       "together when you integrate.", "Sign-out seam:")
bullet(" one analysis costs 100 credits (the ANALYSIS_COST constant). Charge on "
       "submitAnalysis; a pre-receipt upload failure should throw NETWORK_ERROR "
       "and NOT charge.", "Credits:")
bullet(" the S8 screen polls every ~1s and has its own client-side timeout that "
       "maps to PROCESSING_TIMEOUT. Typical job 30–90s is fine.", "Polling:")

# ============================ 8. FILE MAP ===================================
h1("8. Where things live")
table(
    ["Path", "What"],
    [
        ["SPEC.md", "Product spec — the source of truth (14 screens, design)"],
        ["BACKEND.md", "Full API contract + suggested HTTP route mapping"],
        ["IMPROVEMENTS.md", "Agent-team review results + items awaiting approval"],
        ["src/services/api.ts", "The ShuttleCoachApi interface + ApiError"],
        ["src/services/real/realApi.ts", "YOUR file — implement here"],
        ["src/services/mock/", "Reference mock implementation + job simulation"],
        ["src/types/index.ts", "All shared TypeScript types"],
        ["src/constants/", "Pro players, packages, drills, scoring weights"],
        ["src/store/", "Zustand stores (UI cache over the service layer)"],
        ["app/", "The 14 screens (Expo Router)"],
    ],
    widths=[2.5, 3.5],
)

# ============================ 9. REVIEW PASS ================================
h1("9. Quality status")
body(
    "The codebase passed a four-agent review (UX, code, accessibility/perf, and "
    "market-trend). High-priority, low-risk fixes were applied: i18n hardcoded-"
    "string fixes, accessibility (touch targets, labels, contrast), a versioned "
    "mock-DB, an O(n) history-delta computation, store hydration guards, and "
    "more. The full list — including larger items deferred for product "
    "approval — is in IMPROVEMENTS.md."
)
body("Verified before handoff:", bold=True)
bullet(" TypeScript strict (`tsc --noEmit`) — 0 errors.")
bullet(" Web production export (`npx expo export --platform web`) — succeeds.")
bullet(" Web dev bundle — compiles clean (3300+ modules, no errors).")

# ============================ 10. NEXT STEPS ================================
h1("10. Suggested integration order")
for i, step in enumerate([
    "Read SPEC.md (§4 screens, §8 AI analysis) and BACKEND.md.",
    "Stand up auth + getUserProfile / updateUserProfile (incl. the welcome gift).",
    "Wire credits: getCreditBalance, getTransactions, getCreditPackages, "
    "purchaseCredits, refundCredits.",
    "Wire the analysis lifecycle: submitAnalysis → getAnalysisStatus → "
    "getAnalysisResult. This is the CV/AI core.",
    "Wire listAnalyses / getAnalysis / deleteAnalysis / getDrill.",
    "Flip EXPO_PUBLIC_USE_MOCK_API=false, click through every screen, fix any "
    "shape mismatches against src/types.",
]):
    p = doc.add_paragraph(style="List Number")
    p.add_run(step)

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run(
    "Questions on a specific method or type — start from src/services/api.ts "
    "and src/types/index.ts; they are heavily commented. Good luck, Chien!"
)
fr.italic = True
fr.font.color.rgb = SLATE

out = "ShuttleCoach_Integration_Handoff.docx"
doc.save(out)
print("Saved", out)
